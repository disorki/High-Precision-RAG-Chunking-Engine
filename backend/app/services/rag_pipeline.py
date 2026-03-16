import os
import re
import asyncio
import logging
import uuid
from typing import List, Optional, Callable
import httpx
from pypdf import PdfReader
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from openpyxl import load_workbook
from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.config import get_settings

settings = get_settings()
logger = logging.getLogger(__name__)

# поддерживаемые расширения
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.txt'}


def get_file_extension(file_path: str) -> str:
    # получение расширения в нижнем регистре
    ext = os.path.splitext(file_path)[1].lower()
    return ext


class RAGPipeline:
    # обработка документов и генерация эмбеддингов

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=[
                "\n\n",      # абзац (приоритет)
                "\n---\n",   # разделитель таблиц
                "\n| ",      # начало строки таблицы
                "\n",        # перевод строки
                ". ",        # точка
                " ",         # пробел
                ""           # посимвольно (крайний случай)
            ]
        )
        self.ollama_url = settings.ollama_base_url
        self.embedding_model = settings.embedding_model
        self.retry_count = settings.embedding_retry_count
        self.retry_delay = settings.embedding_retry_delay

    # проверка состояния ollama
    async def check_ollama_health(self) -> bool:
        # проверка доступности сервиса
        try:
            async with httpx.AsyncClient(timeout=10.0) as client:
                response = await client.get(f"{self.ollama_url}/api/version")
                response.raise_for_status()
                version_info = response.json()
                logger.info(f"Ollama доступна: версия {version_info.get('version', 'unknown')}")
                return True
        except httpx.ConnectError:
            raise ConnectionError(
                f"Нет связи с Ollama по адресу {self.ollama_url}. Проверьте запуск сервиса."
            )
        except httpx.TimeoutException:
            raise ConnectionError(
                f"Ollama не отвечает (таймаут). Попробуйте позже."
            )
        except Exception as e:
            raise ConnectionError(f"Ошибка проверки Ollama: {e}")

    async def ensure_model_available(self, model_name: str) -> bool:
        # проверка наличия модели
        try:
            async with httpx.AsyncClient(timeout=15.0) as client:
                response = await client.post(
                    f"{self.ollama_url}/api/show",
                    json={"name": model_name}
                )
                if response.status_code == 200:
                    logger.info(f"Модель '{model_name}' доступна")
                    return True
                else:
                    raise ValueError(
                        f"Модель '{model_name}' не загружена в Ollama. Выполните: ollama pull {model_name}"
                    )
        except httpx.ConnectError:
            raise ConnectionError(
                f"Нет связи с Ollama для проверки модели '{model_name}'"
            )
        except ValueError:
            raise
        except Exception as e:
            logger.warning(f"Не удалось проверить модель '{model_name}': {e}")
            return True

    # извлечение текста из pdf
    def extract_text_from_pdf(self, file_path: str) -> List[dict]:
        # извлечение текста по страницам
        try:
            reader = PdfReader(file_path)
            pages = []

            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append({
                        "text": text.strip(),
                        "page_number": i + 1
                    })

            logger.info(f"Извлечено {len(pages)} страниц из PDF: {file_path}")
            return pages
        except Exception as e:
            logger.error(f"Ошибка извлечения текста PDF: {e}")
            raise

    def get_pdf_page_count(self, file_path: str) -> int:
        # получение количества страниц pdf
        try:
            reader = PdfReader(file_path)
            return len(reader.pages)
        except Exception as e:
            logger.error(f"Ошибка получения количества страниц PDF: {e}")
            return 0

    # извлечение текста из docx
    def _iter_block_items(self, doc: DocxDocument):
        # итерация по параграфам и таблицам в порядке следования
        body = doc.element.body
        for child in body:
            if child.tag == qn('w:p'):
                yield Paragraph(child, body)
            elif child.tag == qn('w:tbl'):
                yield DocxTable(child, body)

    def _table_to_text(self, table: DocxTable) -> str:
        # конвертация таблицы docx в текст
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        return "\n".join(rows_text)

    def extract_text_from_docx(self, file_path: str) -> List[dict]:
        # извлечение контента из word с делением на страницы по 3000 симв.
        try:
            doc = DocxDocument(file_path)
            pages = []
            current_text_parts = []
            page_number = 1
            char_count = 0

            for block in self._iter_block_items(doc):
                if isinstance(block, Paragraph):
                    text = block.text.strip()
                    if not text:
                        continue
                    current_text_parts.append(text)
                    char_count += len(text)

                elif isinstance(block, DocxTable):
                    table_text = self._table_to_text(block)
                    if table_text.strip():
                        current_text_parts.append(f"\n[Таблица]\n{table_text}\n")
                        char_count += len(table_text)

                # деление на виртуальные страницы
                if char_count >= 3000:
                    pages.append({
                        "text": "\n".join(current_text_parts),
                        "page_number": page_number
                    })
                    current_text_parts = []
                    char_count = 0
                    page_number += 1

            # остаток
            if current_text_parts:
                pages.append({
                    "text": "\n".join(current_text_parts),
                    "page_number": page_number
                })

            logger.info(f"Извлечено {len(pages)} секций из DOCX: {file_path}")
            return pages
        except Exception as e:
            logger.error(f"Ошибка извлечения текста DOCX: {e}")
            raise

    # извлечение текста из xlsx
    def extract_text_from_xlsx(self, file_path: str) -> List[dict]:
        # каждый лист становится отдельной страницей
        try:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            pages = []

            for sheet_index, sheet_name in enumerate(wb.sheetnames):
                ws = wb[sheet_name]
                rows_text = []

                for row in ws.iter_rows(values_only=True):
                    cell_values = []
                    for cell in row:
                        if cell is not None:
                            cell_values.append(str(cell).strip())
                        else:
                            cell_values.append("")

                    if any(v for v in cell_values):
                        rows_text.append(" | ".join(cell_values))

                if rows_text:
                    sheet_text = f"[Лист: {sheet_name}]\n" + "\n".join(rows_text)
                    pages.append({
                        "text": sheet_text,
                        "page_number": sheet_index + 1
                    })

            wb.close()
            logger.info(f"Извлечено {len(pages)} листов из XLSX: {file_path}")
            return pages
        except Exception as e:
            logger.error(f"Ошибка извлечения текста XLSX: {e}")
            raise

    # извлечение текста из txt
    def extract_text_from_txt(self, file_path: str) -> List[dict]:
        # чтение текстового файла с делением на блоки
        try:
            encodings = ['utf-8', 'cp1251', 'latin-1']
            content = None
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        content = f.read()
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue

            if content is None:
                raise ValueError("Не удалось определить кодировку файла")

            if not content.strip():
                return []

            pages = []
            page_size = 3000
            paragraphs = content.split('\n\n')
            current_parts = []
            char_count = 0
            page_number = 1

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                current_parts.append(para)
                char_count += len(para)

                if char_count >= page_size:
                    pages.append({
                        "text": "\n\n".join(current_parts),
                        "page_number": page_number
                    })
                    current_parts = []
                    char_count = 0
                    page_number += 1

            if current_parts:
                pages.append({
                    "text": "\n\n".join(current_parts),
                    "page_number": page_number
                })

            logger.info(f"Извлечено {len(pages)} секций из TXT: {file_path}")
            return pages
        except Exception as e:
            logger.error(f"Ошибка извлечения текста TXT: {e}")
            raise

    # универсальный метод извлечения
    def extract_text(self, file_path: str) -> List[dict]:
        # автоматическое определение формата по расширению
        ext = get_file_extension(file_path)

        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.xlsx':
            return self.extract_text_from_xlsx(file_path)
        elif ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат: {ext}")

    def get_page_count(self, file_path: str) -> int:
        # получение кол-ва страниц/секций для разных типов файлов
        ext = get_file_extension(file_path)

        if ext == '.pdf':
            return self.get_pdf_page_count(file_path)
        elif ext == '.docx':
            try:
                doc = DocxDocument(file_path)
                para_count = len([p for p in doc.paragraphs if p.text.strip()])
                return max(1, para_count // 30 + (1 if para_count % 30 else 0))
            except Exception:
                return 1
        elif ext == '.xlsx':
            try:
                wb = load_workbook(file_path, read_only=True)
                count = len(wb.sheetnames)
                wb.close()
                return count
            except Exception:
                return 1
        elif ext == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return max(1, len(content) // 3000 + 1)
            except Exception:
                return 1
        else:
            return 0

    def chunk_text(self, pages: List[dict], filename: str = None) -> List[dict]:
        # деление текста на чанки с сохранением метаданных
        chunks = []
        chunk_index = 0

        for page in pages:
            page_text = page["text"].strip()
            if not page_text:
                continue

            page_chunks = self.text_splitter.split_text(page_text)

            for chunk_text in page_chunks:
                chunk_text = chunk_text.strip()
                if not chunk_text:
                    continue
                chunks.append({
                    "chunk_uuid": str(uuid.uuid4()),
                    "text": chunk_text,
                    "page_number": page["page_number"],
                    "chunk_index": chunk_index,
                    "source_filename": filename
                })
                chunk_index += 1

        logger.info(f"Создано {len(chunks)} чанков из {len(pages)} страниц")
        return chunks

    async def generate_context_header(self, text: str) -> Optional[str]:
        # генерация заголовка чанка через llm
        if not settings.enable_context_headers:
            return None

        prompt = (
            f"Write a very brief (15-20 words) descriptive context header "
            f"summarizing the main topic of this text chunk. Output ONLY the header.\n\n"
            f"Chunk:\n{text}\n\nHeader:"
        )

        last_error = None
        for attempt in range(self.retry_count):
            try:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    response = await client.post(
                        f"{self.ollama_url}/api/generate",
                        json={
                            "model": settings.chat_model,
                            "prompt": prompt,
                            "stream": False,
                            "options": {"temperature": 0.1, "num_ctx": 4096}
                        }
                    )
                    response.raise_for_status()
                    data = response.json()
                    header = data.get("response", "").strip()
                    if not header:
                        raise ValueError("Пустой ответ от Ollama.")
                    return header

            except (httpx.ConnectError, httpx.TimeoutException) as e:
                last_error = e
                wait_time = self.retry_delay * (2 ** attempt)
                logger.warning(
                    f"Попытка {attempt + 1}/{self.retry_count} не удалась. Повтор через {wait_time}с..."
                )
                await asyncio.sleep(wait_time)
            except httpx.HTTPStatusError as e:
                if e.response.status_code == 404:
                    raise ValueError(f"Модель '{settings.chat_model}' не найдена в Ollama.")
                last_error = e
                wait_time = self.retry_delay * (2 ** attempt)
                await asyncio.sleep(wait_time)
            except Exception as e:
                last_error = e
                logger.error(f"Ошибка генерации заголовка: {e}")
                raise

        raise ConnectionError(f"Ошибка генерации заголовка. Последняя ошибка: {last_error}")

    async def generate_context_headers_batch(
        self,
        texts: List[str],
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> List[Optional[str]]:
        # пакетная генерация заголовков
        if not settings.enable_context_headers:
            return [None] * len(texts)

        headers: List[Optional[str]] = [None] * len(texts)
        state = {"failed": 0, "completed": 0}
        
        sem = asyncio.Semaphore(settings.embedding_concurrency)

        async def process_one(idx: int, text: str):
            async with sem:
                try:
                    header = await self.generate_context_header(text)
                    headers[idx] = header
                except Exception as e:
                    logger.error(f"Не удалось сгенерировать заголовок для чанка {idx}: {e}")
                    state["failed"] += 1
                
                state["completed"] += 1
                if progress_callback and state["completed"] % 5 == 0:
                    progress_callback(state["completed"], len(texts))

        tasks = [process_one(i, txt) for i, txt in enumerate(texts)]
        await asyncio.gather(*tasks)

        if state["failed"] > 0:
            logger.warning(f"Генерация заголовков: {state['failed']}/{len(texts)} не удалась")
        else:
            logger.info(f"Успешно сгенерировано {len(texts)} заголовков")

        return headers

    async def generate_embedding(self, text: str) -> Optional[List[float]]:
        # генерация эмбеддинга через ollama с повторами
        last_error = None

        for attempt in range(self.retry_count):
            try:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    response = await client.post(
                        f"{self.ollama_url}/api/embeddings",
                        json={
                            "model": self.embedding_model,
                            "prompt": text
                        }
                    )
                    response.raise_for_status()
                    data = response.json()
                    embedding = data.get("embedding")
                    if embedding is None:
                        raise ValueError(f"Нет данных эмбеддинга. Ответ: {data}")
                    return embedding

            except (httpx.ConnectError, httpx.TimeoutException) as e:
                last_error = e
                wait_time = self.retry_delay * (2 ** attempt)
                await asyncio.sleep(wait_time)

            except httpx.HTTPStatusError as e:
                if e.response.status_code == 404:
                    raise ValueError(
                        f"Модель '{self.embedding_model}' не найдена в Ollama. Загрузите ее: ollama pull {self.embedding_model}"
                    )
                last_error = e
                wait_time = self.retry_delay * (2 ** attempt)
                await asyncio.sleep(wait_time)

            except Exception as e:
                last_error = e
                logger.error(f"Ошибка генерации эмбеддинга: {e}")
                raise

        raise ConnectionError(
            f"Ошибка генерации эмбеддинга после {self.retry_count} попыток. Последняя ошибка: {last_error}"
        )

    async def generate_embeddings_batch(
        self,
        texts: List[str],
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> List[Optional[List[float]]]:
        # параллельная генерация эмбеддингов
        embeddings: List[Optional[List[float]]] = [None] * len(texts)
        state = {"failed": 0, "completed": 0}
        
        sem = asyncio.Semaphore(settings.embedding_concurrency)

        async def process_one(idx: int, text: str):
            async with sem:
                try:
                    embedding = await self.generate_embedding(text)
                    embeddings[idx] = embedding
                except Exception as e:
                    logger.error(
                        f"Не удалось сгенерировать эмбеддинг для чанка {idx} "
                        f"(предпросмотр текста: '{text[:80]}...'): {e}"
                    )
                    state["failed"] += 1
                
                state["completed"] += 1
                if progress_callback and state["completed"] % 5 == 0:
                    progress_callback(state["completed"], len(texts))
                if state["completed"] % 10 == 0:
                    logger.info(f"Сгенерировано {state['completed']}/{len(texts)} эмбеддингов ({state['failed']} не удалось)")

        tasks = [process_one(i, txt) for i, txt in enumerate(texts)]
        await asyncio.gather(*tasks)

        if state["failed"] > 0:
            logger.warning(f"Генерация эмбеддингов: {state['failed']}/{len(texts)} не удалась")

        return embeddings

    async def process_document(self, file_path: str) -> List[dict]:
        # полный цикл: проверка, извлечение, чанкинг, эмбеддинги
        await self.check_ollama_health()
        await self.ensure_model_available(self.embedding_model)

        # извлечение текста
        pages = self.extract_text(file_path)
        if not pages:
            raise ValueError("Не удалось извлечь текст из документа")

        # чанкинг
        chunks = self.chunk_text(pages)
        if not chunks:
            raise ValueError("Не удалось создать чанки из текста")

        # эмбеддинги
        texts = [chunk["text"] for chunk in chunks]
        embeddings = await self.generate_embeddings_batch(texts)

        # сборка результата
        processed_chunks = []
        for chunk, embedding in zip(chunks, embeddings):
            if embedding is not None:
                processed_chunks.append({**chunk, "embedding": embedding})
            else:
                logger.warning(f"Пропуск чанка {chunk['chunk_index']} - нет эмбеддинга")

        logger.info(f"Успешно обработано {len(processed_chunks)} чанков")
        return processed_chunks


# синглтон пайплайна
rag_pipeline = RAGPipeline()
